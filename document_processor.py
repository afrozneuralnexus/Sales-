"""
DocumentProcessor – ingests Excel, CSV, PDF, and Word files
and returns a flat list of text chunks with metadata.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import List

import pandas as pd


@dataclass
class DocumentChunk:
    text: str
    source: str
    page: int = 0
    sheet: str = ""
    metadata: dict = field(default_factory=dict)


class DocumentProcessor:
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    # ── Public entry point ────────────────────────────────────────────────

    def process_file(self, uploaded_file) -> List[DocumentChunk]:
        ext = uploaded_file.name.split(".")[-1].lower()
        raw = uploaded_file.read()
        uploaded_file.seek(0)

        if ext in ("xlsx", "xls"):
            return self._process_excel(raw, uploaded_file.name)
        if ext == "csv":
            return self._process_csv(raw, uploaded_file.name)
        if ext == "pdf":
            return self._process_pdf(raw, uploaded_file.name)
        if ext in ("docx", "doc"):
            return self._process_word(raw, uploaded_file.name)
        raise ValueError(f"Unsupported file type: .{ext}")

    # ── File handlers ─────────────────────────────────────────────────────

    def _process_excel(self, raw: bytes, name: str) -> List[DocumentChunk]:
        chunks: List[DocumentChunk] = []
        xl = pd.ExcelFile(io.BytesIO(raw))
        for sheet in xl.sheet_names:
            df = xl.parse(sheet)
            df = df.dropna(how="all").fillna("")
            text = self._dataframe_to_text(df, sheet_name=sheet)
            for chunk in self._split_text(text):
                chunks.append(DocumentChunk(text=chunk, source=name, sheet=sheet))
        return chunks

    def _process_csv(self, raw: bytes, name: str) -> List[DocumentChunk]:
        chunks: List[DocumentChunk] = []
        # Try common encodings
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                df = pd.read_csv(io.BytesIO(raw), encoding=enc)
                break
            except Exception:
                continue
        else:
            raise ValueError("Could not decode CSV file.")

        df = df.dropna(how="all").fillna("")
        text = self._dataframe_to_text(df, sheet_name="CSV")
        for chunk in self._split_text(text):
            chunks.append(DocumentChunk(text=chunk, source=name, sheet="CSV"))
        return chunks

    def _process_pdf(self, raw: bytes, name: str) -> List[DocumentChunk]:
        try:
            import pypdf
        except ImportError:
            raise ImportError("pypdf is required for PDF processing. Run: pip install pypdf")

        chunks: List[DocumentChunk] = []
        reader = pypdf.PdfReader(io.BytesIO(raw))
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            text = re.sub(r"\s+", " ", text).strip()
            if not text:
                continue
            for chunk in self._split_text(text):
                chunks.append(DocumentChunk(text=chunk, source=name, page=page_num))
        return chunks

    def _process_word(self, raw: bytes, name: str) -> List[DocumentChunk]:
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required for Word processing. Run: pip install python-docx")

        doc = docx.Document(io.BytesIO(raw))
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        # Also include tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append(" | ".join(cell.text.strip() for cell in row.cells))
            full_text += "\n" + "\n".join(rows)

        chunks: List[DocumentChunk] = []
        for chunk in self._split_text(full_text):
            chunks.append(DocumentChunk(text=chunk, source=name))
        return chunks

    # ── Helpers ───────────────────────────────────────────────────────────

    def _dataframe_to_text(self, df: pd.DataFrame, sheet_name: str = "") -> str:
        lines = []
        if sheet_name:
            lines.append(f"[Sheet: {sheet_name}]")
        lines.append(f"Columns: {', '.join(str(c) for c in df.columns)}")
        lines.append(f"Total rows: {len(df)}")
        lines.append("")

        # Summary statistics for numeric columns
        num_cols = df.select_dtypes(include="number").columns.tolist()
        if num_cols:
            lines.append("=== Numeric Summary ===")
            for col in num_cols[:20]:  # cap at 20 cols
                s = df[col]
                lines.append(
                    f"{col}: min={s.min():.2f}, max={s.max():.2f}, "
                    f"mean={s.mean():.2f}, sum={s.sum():.2f}, count={s.count()}"
                )
            lines.append("")

        # Raw rows (cap at 500 for large files)
        lines.append("=== Data Rows ===")
        sample = df.head(500)
        for _, row in sample.iterrows():
            lines.append(" | ".join(f"{col}: {val}" for col, val in row.items() if str(val).strip()))

        return "\n".join(lines)

    def _split_text(self, text: str) -> List[str]:
        words = text.split()
        if not words:
            return []
        chunks = []
        start = 0
        while start < len(words):
            end = start + self.chunk_size
            chunk = " ".join(words[start:end])
            chunks.append(chunk)
            start += self.chunk_size - self.chunk_overlap
        return [c for c in chunks if c.strip()]
